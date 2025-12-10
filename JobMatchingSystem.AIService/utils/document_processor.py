import tempfile
import os
from typing import Optional, Tuple
from docx import Document
import PyPDF2
from config.config import Config


class DocumentProcessor:
    """Utility class for processing multiple document types (PDF, DOCX)"""
    
    SUPPORTED_EXTENSIONS = {
        'pdf': ['pdf'],
        'document': ['docx', 'doc']
    }
    
    MAX_FILE_SIZE_MB = 10  # 10MB limit
    
    @classmethod
    def get_file_type(cls, filename: str) -> Optional[str]:
        """Determine file type from extension"""
        if not filename or '.' not in filename:
            return None
            
        ext = filename.lower().split('.')[-1]
        
        for file_type, extensions in cls.SUPPORTED_EXTENSIONS.items():
            if ext in extensions:
                return file_type
        
        return None
    
    @classmethod
    def is_supported_file(cls, filename: str) -> bool:
        """Check if file type is supported"""
        return cls.get_file_type(filename) is not None
    
    @classmethod
    def validate_file(cls, file_content: bytes, filename: str) -> Tuple[bool, Optional[str]]:
        """Validate file size, type, and basic content"""
        
        # Check if file type is supported
        if not cls.is_supported_file(filename):
            return False, "Unsupported file type. Supported formats: PDF, DOCX"
        
        # Check file size
        file_size_mb = len(file_content) / (1024 * 1024)
        if file_size_mb > cls.MAX_FILE_SIZE_MB:
            return False, f"File size ({file_size_mb:.1f}MB) exceeds limit ({cls.MAX_FILE_SIZE_MB}MB)"
        
        # Check if file is empty
        if len(file_content) == 0:
            return False, "File is empty"
        
        return True, None
    
    @classmethod
    def extract_text_from_file(cls, file_content: bytes, filename: str) -> str:
        """Extract text from various file types"""
        file_type = cls.get_file_type(filename)
        
        try:
            if file_type == 'pdf':
                return cls._extract_text_from_pdf(file_content)
            elif file_type == 'document':
                return cls._extract_text_from_docx(file_content)
            else:
                return ""
        except Exception as e:
            print(f"Error extracting text from {filename}: {e}")
            return ""
    
    @classmethod
    def _extract_text_from_pdf(cls, pdf_bytes: bytes) -> str:
        """Extract text from PDF bytes"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(pdf_bytes)
                tmp_path = tmp.name

            reader = PyPDF2.PdfReader(tmp_path)
            text = ""
            
            for page in reader.pages:
                try:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
                except Exception as e:
                    print(f"Error extracting text from page: {e}")
                    continue
            
            # Clean up temp file
            os.unlink(tmp_path)
            
            return text.strip()
            
        except Exception as e:
            print(f"Error processing PDF: {e}")
            return ""
    
    @classmethod
    def _extract_text_from_docx(cls, docx_bytes: bytes) -> str:
        """Extract text from Word document"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(docx_bytes)
                tmp_path = tmp.name
            
            # Read Word document
            doc = Document(tmp_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            # Clean up temp file
            os.unlink(tmp_path)
            
            return text.strip()
            
        except Exception as e:
            print(f"Error processing Word document: {e}")
            return ""
    
    @classmethod
    def get_file_info(cls, file_content: bytes, filename: str) -> dict:
        """Get basic information about the file"""
        try:
            file_type = cls.get_file_type(filename)
            
            info = {
                "filename": filename,
                "file_type": file_type,
                "file_size_mb": len(file_content) / (1024 * 1024),
                "text_length": 0,
                "processing_method": ""
            }
            
            # Extract text and get length
            text = cls.extract_text_from_file(file_content, filename)
            info["text_length"] = len(text)
            
            # Set processing method
            if file_type == 'pdf':
                info["processing_method"] = "PDF text extraction"
                
                # Additional info for PDFs
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                        tmp.write(file_content)
                        tmp_path = tmp.name
                    
                    reader = PyPDF2.PdfReader(tmp_path)
                    info["num_pages"] = len(reader.pages)
                    
                    if reader.metadata:
                        info["has_metadata"] = True
                        info["metadata"] = {
                            "title": reader.metadata.get('/Title', ''),
                            "author": reader.metadata.get('/Author', ''),
                            "creator": reader.metadata.get('/Creator', '')
                        }
                    
                    os.unlink(tmp_path)
                    
                except Exception:
                    pass
                    
            elif file_type == 'document':
                info["processing_method"] = "Word document parsing"
                
                # Additional info for Word docs
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                        tmp.write(file_content)
                        tmp_path = tmp.name
                    
                    doc = Document(tmp_path)
                    info["num_paragraphs"] = len(doc.paragraphs)
                    info["num_tables"] = len(doc.tables)
                    
                    # Check for document properties
                    if doc.core_properties:
                        info["has_metadata"] = True
                        info["metadata"] = {
                            "title": doc.core_properties.title or '',
                            "author": doc.core_properties.author or '',
                            "created": str(doc.core_properties.created) if doc.core_properties.created else ''
                        }
                    
                    os.unlink(tmp_path)
                    
                except Exception:
                    pass
            
            return info
            
        except Exception as e:
            return {
                "filename": filename,
                "file_size_mb": len(file_content) / (1024 * 1024),
                "error": f"Could not analyze file: {str(e)}"
            }